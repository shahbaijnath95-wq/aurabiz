"""
Document Parser - Knowledge Base ke liye files ko text + chunks me convert karta hai.

Supported:
  - PDF  (pypdf, fallback pdfplumber)
  - Word (.docx)  (python-docx)
  - Excel (.xlsx, .xls)  (openpyxl)
  - CSV  (csv)
  - TXT / Markdown  (direct read)

Output: list of text chunks (~CHUNK_SIZE chars, overlap CHUNK_OVERLAP).

Note: image OCR aur PPT abhi supported nahi (future phase).
"""

import os
import csv
from typing import List
from loguru import logger

CHUNK_SIZE = 1200       # ~300-400 tokens
CHUNK_OVERLAP = 200


# ────────────────────────────────────────────────────────────────────
#  Parsers (each returns a single text blob)
# ────────────────────────────────────────────────────────────────────
def _parse_pdf(path: str) -> str:
    text_parts: List[str] = []
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                text_parts.append(t.strip())
    except Exception as e:
        logger.debug("[parser] pypdf failed ({}), trying pdfplumber", e)
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        text_parts.append(t.strip())
        except Exception as e2:
            logger.warning("[parser] pdfplumber also failed: {}", e2)
    return "\n\n".join(text_parts)


def _parse_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.warning("[parser] docx failed: {}", e)
        return ""


def _parse_excel(path: str) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True, read_only=True)
        parts: List[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        logger.warning("[parser] excel failed: {}", e)
        return ""


def _parse_csv(path: str) -> str:
    parts: List[str] = []
    try:
        with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                cells = [c.strip() for c in row if c.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as e:
        logger.warning("[parser] csv failed: {}", e)
    return "\n".join(parts)


def _parse_text(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.warning("[parser] text read failed: {}", e)
        return ""


# ────────────────────────────────────────────────────────────────────
#  Dispatch
# ────────────────────────────────────────────────────────────────────
def parse_file(path: str) -> str:
    """File ko text me convert karo (format auto-detect by extension)."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    if ext in (".docx",):
        return _parse_docx(path)
    if ext in (".xlsx", ".xls"):
        return _parse_excel(path)
    if ext == ".csv":
        return _parse_csv(path)
    if ext in (".txt", ".md", ".markdown", ".json", ".rtf"):
        return _parse_text(path)
    # unknown => try text read
    logger.info("[parser] unknown ext '{}', trying plain text", ext)
    return _parse_text(path)


# ────────────────────────────────────────────────────────────────────
#  Chunking
# ────────────────────────────────────────────────────────────────────
def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Text ko chunks me todo. Sentence-aware: koshish karta hai newline/period pe cut kare
    taaki chunk beech-mein word ka na toote.
    """
    text = (text or "").strip()
    if not text:
        return []

    # Prefer splitting on paragraphs first (better semantic units)
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]

    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 1 <= chunk_size:
            current = (current + "\n" + para).strip() if current else para
        else:
            if current:
                chunks.append(current)
            # long paragraph => hard split
            if len(para) > chunk_size:
                for i in range(0, len(para), chunk_size - overlap):
                    chunks.append(para[i : i + chunk_size])
            current = para
    if current:
        chunks.append(current)

    # Filter too-small tail chunks
    return [c for c in chunks if len(c.strip()) >= 20]
